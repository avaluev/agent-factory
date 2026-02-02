"""Document loaders for various file types."""

import csv
from pathlib import Path
from dataclasses import dataclass
from typing import Any


@dataclass
class Document:
    """Loaded document with content and metadata."""
    content: str
    metadata: dict[str, Any]


class DocumentLoader:
    """Load documents from various file types."""
    
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".pdf", ".docx"}
    
    @classmethod
    def load(cls, path: str | Path) -> Document:
        """Load a document from file path."""
        path = Path(path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        ext = path.suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")
        
        metadata = {
            "source": str(path),
            "filename": path.name,
            "extension": ext
        }
        
        if ext == ".txt":
            content = cls._load_text(path)
        elif ext == ".md":
            content = cls._load_markdown(path)
        elif ext == ".csv":
            content = cls._load_csv(path)
        elif ext == ".pdf":
            content = cls._load_pdf(path)
        elif ext == ".docx":
            content = cls._load_docx(path)
        else:
            content = cls._load_text(path)
        
        return Document(content=content, metadata=metadata)
    
    @classmethod
    def load_directory(
        cls,
        directory: str | Path,
        recursive: bool = True
    ) -> list[Document]:
        """Load all supported documents from a directory."""
        directory = Path(directory)
        documents = []
        
        pattern = "**/*" if recursive else "*"
        for path in directory.glob(pattern):
            if path.is_file() and path.suffix.lower() in cls.SUPPORTED_EXTENSIONS:
                try:
                    documents.append(cls.load(path))
                except Exception as e:
                    # Log error but continue with other files
                    print(f"Warning: Failed to load {path}: {e}")
        
        return documents
    
    @staticmethod
    def _load_text(path: Path) -> str:
        """Load plain text file."""
        return path.read_text(encoding="utf-8")
    
    @staticmethod
    def _load_markdown(path: Path) -> str:
        """Load markdown file (treated as text)."""
        return path.read_text(encoding="utf-8")
    
    @staticmethod
    def _load_csv(path: Path) -> str:
        """Load CSV file and convert to text."""
        rows = []
        with open(path, newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                row_text = " | ".join(f"{k}: {v}" for k, v in row.items())
                rows.append(row_text)
        return "\n".join(rows)
    
    @staticmethod
    def _load_pdf(path: Path) -> str:
        """Load PDF file."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("pdfplumber is required for PDF support: pip install pdfplumber")
    
    @staticmethod
    def _load_docx(path: Path) -> str:
        """Load DOCX file."""
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            raise ImportError("python-docx is required for DOCX support: pip install python-docx")
